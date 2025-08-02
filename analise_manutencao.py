import os
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches

# ========== CONFIGURAÇÕES INICIAIS ==========
sns.set(style="whitegrid")
os.makedirs("imagens", exist_ok=True)

# ========== LEITURA DOS DADOS ==========
df = pd.read_csv("manutencao_turbinas.csv", parse_dates=["Data"])
df["AnoMes"] = df["Data"].dt.to_period("M").astype(str)

# ========== ANÁLISE 1: MTTR por Equipamento ==========
mttr = df.groupby("Equipamento")["Tempo_Parado_Horas"].mean().sort_values()

plt.figure(figsize=(8, 4))
sns.barplot(x=mttr.values, y=mttr.index, palette="Blues_r")
plt.title("MTTR por Equipamento")
plt.xlabel("Tempo Médio Parado (Horas)")
plt.tight_layout()
plt.savefig("imagens/grafico_mttr.png")
plt.close()

# ========== ANÁLISE 2: Custo por Tipo de Manutenção ==========
custo_tipo = df.groupby("Tipo")["Custo"].sum().sort_values()

plt.figure(figsize=(6, 6))
custo_tipo.plot(kind="pie", autopct="%1.1f%%", startangle=90, colormap="Pastel1")
plt.title("Distribuição de Custo por Tipo de Manutenção")
plt.ylabel("")
plt.tight_layout()
plt.savefig("imagens/grafico_custo.png")
plt.close()

# ========== ANÁLISE 3: Evolução Mensal do MTTR ==========
evolucao_mttr = df.groupby(["AnoMes", "Equipamento"])["Tempo_Parado_Horas"].mean().reset_index()

plt.figure(figsize=(10, 6))
for eqp in evolucao_mttr["Equipamento"].unique():
    sub = evolucao_mttr[evolucao_mttr["Equipamento"] == eqp]
    plt.plot(sub["AnoMes"], sub["Tempo_Parado_Horas"], marker="o", label=eqp)
plt.title("Evolução Mensal do MTTR por Equipamento")
plt.xlabel("Ano-Mês")
plt.ylabel("MTTR (Horas)")
plt.legend()
plt.xticks(rotation=45)
plt.tight_layout()
plt.savefig("imagens/grafico_evolucao_mttr.png")
plt.close()

# ========== ANÁLISE 4: Top 5 Equipamentos por Custo ==========
top5 = df.groupby("Equipamento")["Custo"].sum().nlargest(5).sort_values()

plt.figure(figsize=(8, 4))
sns.barplot(x=top5.values, y=top5.index, palette="Reds_r")
plt.title("Top 5 Equipamentos por Custo Total")
plt.xlabel("Custo Total (R$)")
plt.tight_layout()
plt.savefig("imagens/grafico_top5_custo.png")
plt.close()

# ========== ANÁLISE 5: Correlação Tempo x Custo ==========
plt.figure(figsize=(6, 4))
sns.scatterplot(data=df, x="Tempo_Parado_Horas", y="Custo", hue="Tipo")
plt.title("Correlação entre Tempo Parado e Custo")
plt.tight_layout()
plt.savefig("imagens/grafico_correlacao.png")
plt.close()

# ========== RELATÓRIO TÉCNICO ==========
doc = Document()
doc.add_heading("Relatório Técnico de Eficiência da Manutenção - Parque Eólico", 0)

doc.add_paragraph(
    "Este relatório apresenta uma análise dos dados de manutenção realizados em turbinas eólicas. "
    "As análises visam avaliar o desempenho dos equipamentos, os custos associados às intervenções "
    "e possíveis correlações entre o tempo de parada e o custo das manutenções."
)

doc.add_heading("1. MTTR por Equipamento", level=1)
doc.add_paragraph(
    "O gráfico a seguir apresenta o Tempo Médio de Reparo (MTTR) por equipamento. "
    "Esse indicador permite identificar quais turbinas estão levando mais tempo em média para serem restauradas ao funcionamento."
)
doc.add_picture("imagens/grafico_mttr.png", width=Inches(5.5))

doc.add_heading("2. Distribuição de Custo por Tipo de Manutenção", level=1)
doc.add_paragraph(
    "A análise de distribuição de custos mostra a proporção dos valores gastos com manutenções corretivas, preditivas e preventivas. "
    "Essa informação é útil para avaliar o equilíbrio entre estratégias de manutenção."
)
doc.add_picture("imagens/grafico_custo.png", width=Inches(5.0))

doc.add_heading("3. Evolução Mensal do MTTR por Equipamento", level=1)
doc.add_paragraph(
    "Este gráfico mostra a tendência mensal do MTTR ao longo do período analisado. "
    "É possível observar oscilações no desempenho da manutenção ao longo do tempo."
)
doc.add_picture("imagens/grafico_evolucao_mttr.png", width=Inches(6.0))

doc.add_heading("4. Top 5 Equipamentos por Custo Total", level=1)
doc.add_paragraph(
    "Os equipamentos com maior custo acumulado são listados a seguir. "
    "Essa informação pode auxiliar na priorização de ações de melhoria ou substituição."
)
doc.add_picture("imagens/grafico_top5_custo.png", width=Inches(5.5))

doc.add_heading("5. Correlação entre Tempo Parado e Custo", level=1)
doc.add_paragraph(
    "A dispersão entre o tempo de parada e o custo de cada manutenção sugere uma relação entre a duração da falha e seu impacto financeiro. "
    "Essa visualização é útil para identificar anomalias e padrões de alto custo."
)
doc.add_picture("imagens/grafico_correlacao.png", width=Inches(5.0))

doc.save("relatorio_tecnico_manutencao.docx")